# backend/app/core/export_docx.py
import io
import json
from docx import Document

def build_docx_bytes(project, db=None):
    """
    Minimal docx exporter: builds a document with project title and each item's
    first section content (or raw content). Returns bytes.
    Requires python-docx: pip install python-docx
    """
    doc = Document()
    title = getattr(project, "title", "") or "Project"
    doc.add_heading(title, level=1)

    # iterate items sorted by order if available
    items = getattr(project, "items", []) or []
    items = sorted(items, key=lambda it: (getattr(it, "order", 0) or 0))

    for it in items:
        doc.add_heading(getattr(it, "title", "") or "", level=2)
        content_md = ""
        try:
            # item.content may be JSON string with sections
            raw = getattr(it, "content", "") or ""
            parsed = json.loads(raw) if isinstance(raw, str) and raw.strip().startswith("{") else None
            if parsed and parsed.get("sections"):
                content_md = parsed["sections"][0].get("content_md", "")
            else:
                content_md = raw or ""
        except Exception:
            content_md = getattr(it, "content", "") or ""
        # docx can't render markdown — put raw markdown/plain text
        if content_md:
            doc.add_paragraph(content_md)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
