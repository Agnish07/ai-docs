from io import BytesIO
from docx import Document

def assemble_docx(project_title, items):
    doc = Document()
    doc.add_heading(project_title, level=1)
    for it in items:
        doc.add_heading(it.get("title",""), level=2)
        doc.add_paragraph(it.get("content",""))
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
